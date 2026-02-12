"""Export corpus et résultats KWIC (TXT, CSV, JSON, JSONL segmenté, Word)."""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document

from howimetyourcorpus.core.models import EpisodeRef
from howimetyourcorpus.core.storage.db import KwicHit
from howimetyourcorpus.core.segment import segment_utterances, segment_utterances_into_phrases


def _corpus_row(ref: EpisodeRef, clean_text: str) -> dict:
    return {
        "episode_id": ref.episode_id,
        "season": ref.season,
        "episode": ref.episode,
        "title": ref.title or "",
        "clean_text": clean_text,
    }


def export_corpus_txt(episodes: list[tuple[EpisodeRef, str]], path: Path) -> None:
    """Exporte le corpus en TXT : une section par épisode (## id - titre + texte)."""
    with path.open("w", encoding="utf-8") as f:
        for ref, text in episodes:
            f.write(f"## {ref.episode_id} - {ref.title or ''}\n\n")
            f.write(text.strip())
            f.write("\n\n")
    return None


def export_corpus_csv(episodes: list[tuple[EpisodeRef, str]], path: Path) -> None:
    """Exporte le corpus en CSV : episode_id, season, episode, title, clean_text."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["episode_id", "season", "episode", "title", "clean_text"])
        for ref, text in episodes:
            w.writerow([ref.episode_id, ref.season, ref.episode, ref.title or "", text])
    return None


def export_corpus_json(episodes: list[tuple[EpisodeRef, str]], path: Path) -> None:
    """Exporte le corpus en JSON : liste d'objets { episode_id, season, episode, title, clean_text }."""
    data = [_corpus_row(ref, text) for ref, text in episodes]
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return None


def export_corpus_docx(episodes: list[tuple[EpisodeRef, str]], path: Path) -> None:
    """Exporte le corpus en Word (.docx) : un titre par épisode, puis le texte en paragraphes."""
    doc = Document()
    doc.add_heading("Corpus exporté", 0)
    for ref, text in episodes:
        doc.add_heading(f"{ref.episode_id} — {ref.title or ''}", level=1)
        for block in (t.strip() for t in text.split("\n\n") if t.strip()):
            doc.add_paragraph(block)
        doc.add_paragraph()
    doc.save(str(path))
    return None


def export_kwic_csv(hits: list[KwicHit], path: Path) -> None:
    """Exporte les résultats KWIC en CSV (inclut segment_id/kind/cue_id/lang si présents)."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        row0 = ["episode_id", "title", "left", "match", "right", "position", "score"]
        if hits and (getattr(hits[0], "segment_id", None) or getattr(hits[0], "kind", None) or getattr(hits[0], "cue_id", None) or getattr(hits[0], "lang", None)):
            if getattr(hits[0], "segment_id", None) or getattr(hits[0], "kind", None):
                row0.extend(["segment_id", "kind"])
            if getattr(hits[0], "cue_id", None) or getattr(hits[0], "lang", None):
                row0.extend(["cue_id", "lang"])
        w.writerow(row0)
        for h in hits:
            r = [h.episode_id, h.title, h.left, h.match, h.right, h.position, h.score]
            if len(row0) > 7:
                r.extend([getattr(h, "segment_id", "") or "", getattr(h, "kind", "") or ""])
            if len(row0) > 9:
                r.extend([getattr(h, "cue_id", "") or "", getattr(h, "lang", "") or ""])
            w.writerow(r)
    return None


def export_kwic_tsv(hits: list[KwicHit], path: Path) -> None:
    """Exporte les résultats KWIC en TSV (inclut segment_id/kind si présents)."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f, delimiter="\t")
        row0 = ["episode_id", "title", "left", "match", "right", "position", "score"]
        if hits and (getattr(hits[0], "segment_id", None) or getattr(hits[0], "kind", None) or getattr(hits[0], "cue_id", None) or getattr(hits[0], "lang", None)):
            if getattr(hits[0], "segment_id", None) or getattr(hits[0], "kind", None):
                row0.extend(["segment_id", "kind"])
            if getattr(hits[0], "cue_id", None) or getattr(hits[0], "lang", None):
                row0.extend(["cue_id", "lang"])
        w.writerow(row0)
        for h in hits:
            r = [h.episode_id, h.title, h.left, h.match, h.right, h.position, h.score]
            if len(row0) > 7:
                r.extend([getattr(h, "segment_id", "") or "", getattr(h, "kind", "") or ""])
            if len(row0) > 9:
                r.extend([getattr(h, "cue_id", "") or "", getattr(h, "lang", "") or ""])
            w.writerow(r)
    return None


def export_kwic_json(hits: list[KwicHit], path: Path) -> None:
    """Exporte les résultats KWIC en JSON (inclut segment_id/kind si présents)."""
    data = []
    for h in hits:
        row = {
            "episode_id": h.episode_id,
            "title": h.title,
            "left": h.left,
            "match": h.match,
            "right": h.right,
            "position": h.position,
            "score": h.score,
        }
        if getattr(h, "segment_id", None):
            row["segment_id"] = h.segment_id
        if getattr(h, "kind", None):
            row["kind"] = h.kind
        if getattr(h, "cue_id", None):
            row["cue_id"] = h.cue_id
        if getattr(h, "lang", None):
            row["lang"] = h.lang
        data.append(row)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return None


def export_kwic_jsonl(hits: list[KwicHit], path: Path) -> None:
    """Exporte les résultats KWIC en JSONL (une ligne JSON par hit)."""
    with path.open("w", encoding="utf-8") as f:
        for h in hits:
            row = {
                "episode_id": h.episode_id,
                "title": h.title,
                "left": h.left,
                "match": h.match,
                "right": h.right,
                "position": h.position,
                "score": h.score,
            }
            if getattr(h, "segment_id", None):
                row["segment_id"] = h.segment_id
            if getattr(h, "kind", None):
                row["kind"] = h.kind
            if getattr(h, "cue_id", None):
                row["cue_id"] = h.cue_id
            if getattr(h, "lang", None):
                row["lang"] = h.lang
            f.write(json.dumps(row, ensure_ascii=False) + "\n")
    return None


# --- Export segmenté (Phase 2 : utterances / phrases, JSONL + CSV) ---


def export_corpus_utterances_jsonl(
    episodes: list[tuple[EpisodeRef, str]], path: Path
) -> None:
    """Exporte le corpus segmenté en utterances : une ligne JSON par utterance (JSONL)."""
    with path.open("w", encoding="utf-8") as f:
        for ref, text in episodes:
            for u in segment_utterances(text):
                obj = {
                    "episode_id": ref.episode_id,
                    "season": ref.season,
                    "episode": ref.episode,
                    "title": ref.title or "",
                    "speaker": u.speaker,
                    "text": u.text,
                    "index": u.index,
                }
                f.write(json.dumps(obj, ensure_ascii=False) + "\n")
    return None


def export_corpus_utterances_csv(
    episodes: list[tuple[EpisodeRef, str]], path: Path
) -> None:
    """Exporte le corpus segmenté en utterances : CSV (episode_id, season, episode, title, speaker, text, index)."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["episode_id", "season", "episode", "title", "speaker", "text", "index"])
        for ref, text in episodes:
            for u in segment_utterances(text):
                w.writerow([
                    ref.episode_id,
                    ref.season,
                    ref.episode,
                    ref.title or "",
                    u.speaker or "",
                    u.text,
                    u.index,
                ])
    return None


def export_corpus_phrases_jsonl(
    episodes: list[tuple[EpisodeRef, str]], path: Path
) -> None:
    """Exporte le corpus segmenté en phrases : une ligne JSON par phrase (JSONL)."""
    with path.open("w", encoding="utf-8") as f:
        for ref, text in episodes:
            for ph in segment_utterances_into_phrases(text):
                obj = {
                    "episode_id": ref.episode_id,
                    "season": ref.season,
                    "episode": ref.episode,
                    "title": ref.title or "",
                    "speaker": ph.speaker,
                    "text": ph.text,
                    "index": ph.index,
                }
                f.write(json.dumps(obj, ensure_ascii=False) + "\n")
    return None


def export_corpus_phrases_csv(
    episodes: list[tuple[EpisodeRef, str]], path: Path
) -> None:
    """Exporte le corpus segmenté en phrases : CSV (episode_id, season, episode, title, speaker, text, index)."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["episode_id", "season", "episode", "title", "speaker", "text", "index"])
        for ref, text in episodes:
            for ph in segment_utterances_into_phrases(text):
                w.writerow([
                    ref.episode_id,
                    ref.season,
                    ref.episode,
                    ref.title or "",
                    ph.speaker or "",
                    ph.text,
                    ph.index,
                ])
    return None
